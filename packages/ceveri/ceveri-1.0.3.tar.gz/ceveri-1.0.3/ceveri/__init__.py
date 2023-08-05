import os
import json
import docx
import time
import pandas as pd
import six
from googletrans import Translator

def cevir(file, save_path, source_language=None):
    allowed_extensions = [".txt", ".csv", ".xlsx", ".json", ".xml", ".pkl", ".docx"]
    translator = Translator()
    extension = os.path.splitext(file)[1]

    def translate_text(target, text, source_language, detect_source_language):
        if isinstance(text, six.binary_type):
            text = text.decode("utf-8")

        if detect_source_language:
            source_language = translator.detect(text).lang
            result = translator.translate(text, src=source_language, dest=target)

        else:
            result = translator.translate(text, src=source_language, dest=target)

        return result.text

    def translate_json(json_file, source_language, detect_source_language):
        search_dict = json.load(open(json_file, "r", encoding="utf-8"))
        def translate_recursively(search_dict, field, detect_source_language):
            for key, value in search_dict.items():
                if isinstance(value, field):
                    if "https://" not in value:
                        translation = translate_text("tr", value, source_language, detect_source_language)
                        search_dict[key] = translation

                elif isinstance(value, dict):
                    results = translate_recursively(value, field, detect_source_language)

                elif isinstance(value, (list,tuple)):
                    for item in value:
                        if isinstance(item, dict):
                            more_results = translate_recursively(item, field, detect_source_language)

                        elif isinstance(item, field):
                            if "https://" not in item:
                                translation = translate_text("tr", item, source_language, detect_source_language)
                                value[value.index(item)] = translation

                    search_dict[key] = value
                        
            return search_dict
        
        translated_dict = translate_recursively(search_dict, str, detect_source_language)
        return translated_dict

    def translate_excel(excel_file, source_language, detect_source_language):
        df = pd.read_excel(excel_file)

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)

        return df

    def translate_csv(csv_file, source_language, detect_source_language):
        df = pd.read_csv(csv_file, error_bad_lines=False)

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)

        return df

    def translate_xml(xml_file, source_language, detect_source_language):
        df = pd.read_xml(xml_file)

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)

        return df

    def translate_pkl(pkl_file, source_language, detect_source_language):
        df = pd.read_pickle(pkl_file)

        for index, row in list(df.iterrows()):
            for column in list(df.columns):
                if isinstance(row[column], str):
                    translation = translate_text("tr", row[column], source_language, detect_source_language)
                    df.replace(row[column], translation, inplace=True)

        return df

    def translate_txt(txt_file, source_language, detect_source_language):
        data = open(txt_file, "r", encoding="utf-8").readlines()
        translations = ["\n"] * len(data)

        for datum in data:
            if datum != "\n":
                translation = translate_text("tr", datum.replace("\n", ""), source_language, detect_source_language)
                translations[data.index(datum)] = translation
            else:
                translations[data.index(datum)] = datum

        return translations

    def translate_docx(docx_file, source_language, detect_source_language):
        doc = docx.Document(docx_file)
        translated_doc = docx.Document()

        paragraphs = doc.paragraphs
        for paragraph in paragraphs:
            translated_paragraph = translate_text(paragraph.text, source_language, detect_source_language)
            translated_doc.add_paragraph(translated_paragraph)

        return translated_doc

    if extension.lower() not in allowed_extensions:
        raise TypeError("Uzantı dosyası desteklenilmiyor. Lütfen desteklenilen uzantıları görmek için https://ceveri.readthedocs.io adresinden ÇeVeri dokümantasyonlarını ziyaret ediniz.")

    if source_language == None:
        detect_source_language = True
    else:
        detect_source_language = False

    if extension.lower() == ".json":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_dict = translate_json(file, source_language, detect_source_language)
        with open(save_path, 'w', encoding="utf-8") as f:
            json.dump(translated_dict, f, indent=4, ensure_ascii=False)
        after = time.time()

    if extension.lower() == ".xlsx":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_df = translate_excel(file, source_language, detect_source_language)
        translated_df.to_excel(save_path, index=False)
        after = time.time()

    if extension.lower() == ".csv":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_df = translate_csv(file, source_language, detect_source_language)
        translated_df.to_csv(save_path, index=False)
        after = time.time()

    if extension.lower() == ".xml":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_df = translate_xml(file, source_language, detect_source_language)
        translated_df.to_xml(save_path, index=False)
        after = time.time()

    if extension.lower() == ".pkl":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_df = translate_pkl(file, source_language, detect_source_language)
        translated_df.to_pickle(save_path, index=False)
        after = time.time()

    if extension.lower() == ".txt":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translation_list = translate_txt(file, source_language, detect_source_language)
        with open(save_path, "w", encoding="utf-8") as f:
            for line in translation_list:
                if line != "\n":
                    f.writelines(line+"\n")
                else:
                    f.writelines(line)
        after = time.time()

    if extension.lower() == ".docx":
        before = time.time()
        print("Çeviri işlemi başladı.")
        translated_docx = translate_docx(file, source_language, detect_source_language)
        translated_docx.save(save_path)
        after = time.time()

    execution_time = after - before

    print("Dosya başarıyla çevrildi.")
    print("İşlem süresi: " + str(execution_time) + " saniye.")
    print("Çevrilen dosyaya " + save_path + " adresinden erişebilirsiniz.")
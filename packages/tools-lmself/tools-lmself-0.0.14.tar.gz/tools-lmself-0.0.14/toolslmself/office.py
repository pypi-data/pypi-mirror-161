import docx  # pip install python-docx -i https://mirrors.aliyun.com/pypi/simple
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 段落对齐方式
from docx.shared import RGBColor  # 设置字体颜色
from docx.shared import Pt, Cm, Inches  # 字体大小，不一定全用到

import openpyxl  # pip install openpyxl -i https://mirrors.aliyun.com/pypi/simple


class docxOP:
    """操作excel文件"""

    @staticmethod
    def read_docx(file):
        D = docx.Document(file)
        TB, TE = [], []
        for table in D.tables:
            temp = []
            for row in table.rows:
                r_s = [cell.text.strip() for cell in row.cells]
                temp.append(r_s)
            TB.append(temp)

        for para in D.paragraphs:
            TE.append(para.text.strip())
        return TB, TE

    class reader:
        def __init__(self, docxfile):
            """
            docx reader
            :param docxfile: docx文件名
            """
            self.D = docx.Document(docxfile)
            self.TE = []
            self.TB = []

            self.get_text()
            self.get_table()

        def __del__(self):
            pass

        def get_text(self):
            """
            获取段落文本
            :return:
            """
            for para in self.D.paragraphs:
                self.TE.append(para.text)

        def get_table(self):
            """
            获取表格
            :return:
            """
            for table in self.D.tables:
                temp = []
                for row in table.rows:
                    r_s = [cell.text.strip() for cell in row.cells]
                    temp.append(r_s)
                self.TB.append(temp)

        def save_picture(self, save_picture_path):
            """
            保存图片到指定路径下
            :param save_picture_path: 保存图片的路径
            :return:
            """
            for order, i in enumerate(self.D.inline_shapes):
                blip = self.D.inline_shapes[order]._inline.graphic.graphicData.pic.blipFill.blip
                image_part = self.D.part.related_parts[blip.embed]
                with open(f'{save_picture_path}/test{order}.png', 'wb') as f:
                    f.write(image_part._blob)

    class writer:
        def __init__(self, docxfile, mode='a'):
            """
            docx writer
            :param docxfile: docx文件名
            :param mode: 写|追加 w|a
            """
            self.docxfile = docxfile
            if mode == 'w':
                self.D = docx.Document()
            elif mode == 'a':
                try:
                    self.D = docx.Document(self.docxfile)
                except docx.opc.exceptions.PackageNotFoundError:
                    self.D = docx.Document()

        def __del__(self):
            self.D.save(self.docxfile)

        def add_heading(self, text, level=1):
            """
            添加标题
            :param text: 标题文本
            :param level: 标题等级
            :return:
            """
            self.D.add_heading(text, level)

        def add_paragraph(self, text=None, align=None, style=None):
            """
            添加段落
            :param text: 段落文本
            :param align: 对齐方式
            :return:
            """
            self.p = self.D.add_paragraph(text)
            self.p.style = style
            if align == 'LEFT':
                self.p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif align == 'CENTER':
                self.p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == 'RIGHT':
                self.p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif align == 'JUSTIFY':
                self.p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif align == 'DISTRIBUTE':
                self.p.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
            return self

        def add_run(self, text, sizePt=10.5, B=False, I=False, U=False, color=(0, 0, 0)):
            """
            添加 块
            :param sizePt: 字体大小Pt
            :param B: 粗体
            :param I: 斜体
            :param U: 下划线
            :param color: 颜色
            :param text: 块文本
            :return:
            """
            run = self.p.add_run(text)
            run.bold = B
            run.italic = I
            run.underline = U
            run.font.size = Pt(sizePt)
            r, g, b = [i for i in color]
            run.font.color.rgb = RGBColor(r, g, b)
            return self

        def add_picture(self, pic_file, width=None, height=None):
            """
            添加图片
            :param pic_file: 图片
            :param width: 宽
            :param height: 高
            :return:
            """
            self.D.add_picture(pic_file, width, height)

        def add_table(self, data, style='Table Grid'):
            """
            添加表格
            :param data: 二维数组
            :param style: 表格样式
            :return:
            """
            table = self.D.add_table(rows=0, cols=len(data[0]), style=style)
            # for order, i in data[0]:
            #     table.rows[0].cells[0].text = str(data[0][0])
            for i in data:
                row_cells = table.add_row().cells
                for order, ii in enumerate(i):
                    row_cells[order].text = str(ii)

        def save_as(self, save_docxfile):
            """
            docx文件另存为
            :param save_docxfile: 另存为文件名
            :return:
            """
            self.docxfile = save_docxfile
            return self

    def markdown2docx(self, markdownfile, docxfile):
        pass


class ExcelOP:
    """操作excel文件"""

    class reader:
        def __init__(self, excelfile):
            """
            excel reader
            :param excelfile: excel文件名
            """
            self.excelfile = excelfile
            self.wb = openpyxl.load_workbook(self.excelfile)

        def __del__(self):
            self.wb.close()  # 这一行加不加关系不大,加了可以释放内存

        def sheet_lists(self):
            """
            excel所有的sheet名
            :return:
            """
            return self.wb.sheetnames

        def sheet_data(self, sheetname):
            """
            读取数据
            :param sheetname:
            :return:
            """
            sh = self.wb[sheetname]
            rows = list(sh.rows)
            data = [['' if r.value is None else r.value for r in row] for row in rows]
            return data

        def sheet(self, sheetname):
            """
            读取cxcel文件一个sheet
            :param sheetname: sheet名
            :return:
            """
            return self.sheet_data(sheetname)

        def sheets(self):
            """
            读取cxcel文件一个所有sheet
            :return:
            """
            data = []
            for i in self.wb.sheetnames:
                data.append([i, self.sheet_data(i)])
            return data

    class writer:
        def __init__(self, excelfile):
            """
            excel writer
            :param excelfile: excel文件名
            """
            self.excelfile = excelfile

            try:
                self.wb = openpyxl.load_workbook(self.excelfile)
            except FileNotFoundError:
                self.wb = openpyxl.Workbook()

        def __del__(self):
            self.wb.close()  # 这一行加不加关系不大,加了可以释放内存

        def sheet_lists(self):
            """
            excel所有的sheet名
            :return:
            """
            return self.wb.sheetnames

        def sheet(self, sheetname, mode):
            """
            设置sheet名
            :param sheetname: sheet名
            :param mode: 写|追加 w|a
            :return:
            """
            if sheetname in self.wb.sheetnames:  # 有 sheet
                if mode == 'w':
                    index = self.wb.sheetnames.index(sheetname)
                    self.wb.remove(self.wb[sheetname])  # 删除 del self.wb[sheetname]
                    self.ws = self.wb.create_sheet(sheetname, index)
                elif mode == 'a':
                    self.ws = self.wb[sheetname]
            else:  # 无 sheet
                self.ws = self.wb.create_sheet(sheetname, 0)
            return self

        def writerow(self, row):
            """
            excel写入一行数据
            :param row: 一维数据
            :return:
            """
            self.ws.append(row)  # 用append插入一行数据
            self.wb.save(self.excelfile)
            return self

        def writerows(self, rows):
            """
            excel写入多行数据
            :param rows: 二维数组
            :return:
            """
            for r in rows:
                self.ws.append(r)  # 用append插入一行数据
            self.wb.save(self.excelfile)
            return self

        def save_as(self, save_excelfile):
            """
            excel文件另存为
            :param save_excelfile: 另存为文件名
            :return:
            """
            self.excelfile = save_excelfile
            return self

        def delete_sheet(self, sheetname):
            """
            删除sheet
            :param sheetname: sheet表
            :return:
            """
            if sheetname in self.wb.sheetnames:
                del self.wb[sheetname]
                self.wb.save(self.excelfile)
            else:
                print(f'{self.excelfile} 文件中没有 {sheetname}.')
            return self
